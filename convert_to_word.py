import re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_styles(doc):
    styles = doc.styles
    
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = '宋体'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    h1_style = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = '黑体'
    h1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    h2_style = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = '黑体'
    h2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(4)
    
    h3_style = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.font.name = '黑体'
    h3_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(4)
    
    h4_style = styles.add_style('CustomH4', WD_STYLE_TYPE.PARAGRAPH)
    h4_style.font.name = '黑体'
    h4_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    h4_style.font.size = Pt(11)
    h4_style.font.bold = True
    h4_style.paragraph_format.space_before = Pt(6)
    h4_style.paragraph_format.space_after = Pt(4)
    
    body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = '宋体'
    body_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    body_style.font.size = Pt(12)
    body_style.paragraph_format.first_line_indent = Cm(0.74)
    body_style.paragraph_format.line_spacing = 1.5
    
    code_style = styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Cm(0.5)
    
    return styles

def parse_table(lines, start_idx):
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line or not line.startswith('|'):
            break
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells and not all(c.replace('-', '').replace(':', '') == '' for c in cells):
            rows.append(cells)
        i += 1
    return rows, i

def add_table(doc, rows):
    if not rows:
        return
    
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def convert_md_to_docx(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    styles = create_styles(doc)
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    
    while i < len(lines):
        line = lines[i]
        
        if line.strip().startswith('```'):
            if in_code_block:
                for code_line in code_content:
                    p = doc.add_paragraph(code_line, style='CustomCode')
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        stripped = line.strip()
        
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title = stripped[2:].strip()
            p = doc.add_paragraph(title, style='CustomTitle')
        
        elif stripped.startswith('## '):
            title = stripped[3:].strip()
            p = doc.add_paragraph(title, style='CustomH1')
        
        elif stripped.startswith('### '):
            title = stripped[4:].strip()
            p = doc.add_paragraph(title, style='CustomH2')
        
        elif stripped.startswith('#### '):
            title = stripped[5:].strip()
            p = doc.add_paragraph(title, style='CustomH3')
        
        elif stripped.startswith('##### '):
            title = stripped[6:].strip()
            p = doc.add_paragraph(title, style='CustomH4')
        
        elif stripped.startswith('|') and '|' in stripped[1:]:
            rows, new_i = parse_table(lines, i)
            add_table(doc, rows)
            i = new_i
            continue
        
        elif stripped.startswith('---'):
            pass
        
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
        
        elif stripped.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            text = stripped[3:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
        
        elif stripped:
            text = stripped
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            
            p = doc.add_paragraph(style='CustomBody')
            run = p.add_run(text)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
        
        i += 1
    
    doc.save(docx_file)
    print(f'转换完成: {docx_file}')

if __name__ == '__main__':
    convert_md_to_docx('毕业设计论文.md', '毕业设计论文.docx')
